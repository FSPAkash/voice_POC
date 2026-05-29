from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parent.parent
SOURCE = ROOT / "PRICING_COMPARISON.md"
OUT = ROOT / "PRICING_COMPARISON.docx"


def clean_inline(text: str) -> str:
    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1 (\2)", text)
    text = text.replace("**", "")
    text = text.replace("`", "")
    return text.strip()


def split_table_row(line: str) -> list[str]:
    return [clean_inline(cell) for cell in line.strip().strip("|").split("|")]


def is_table_divider(line: str) -> bool:
    cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
    return bool(cells) and all(cell and all(ch in "-:" for ch in cell) for cell in cells)


def add_table(doc: Document, lines: list[str]) -> None:
    rows = [split_table_row(line) for line in lines if not is_table_divider(line)]
    if not rows:
        return

    width = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=width)
    table.style = "Light Grid Accent 1"

    for row_index, row in enumerate(rows):
        for col_index in range(width):
            cell = table.rows[row_index].cells[col_index]
            cell.text = row[col_index] if col_index < len(row) else ""
            if row_index == 0:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
    doc.add_paragraph()


def add_paragraph(doc: Document, lines: list[str]) -> None:
    text = clean_inline(" ".join(line.strip() for line in lines))
    if text:
        doc.add_paragraph(text)


def render_docx(markdown: str) -> Document:
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10)

    for section in doc.sections:
        section.left_margin = Inches(0.65)
        section.right_margin = Inches(0.65)
        section.top_margin = Inches(0.65)
        section.bottom_margin = Inches(0.65)

    lines = markdown.splitlines()
    paragraph_buffer: list[str] = []
    table_buffer: list[str] = []

    def flush_paragraph() -> None:
        nonlocal paragraph_buffer
        add_paragraph(doc, paragraph_buffer)
        paragraph_buffer = []

    def flush_table() -> None:
        nonlocal table_buffer
        add_table(doc, table_buffer)
        table_buffer = []

    for raw_line in lines:
        line = raw_line.rstrip()

        if table_buffer and not line.strip().startswith("|"):
            flush_table()

        if line.strip().startswith("|"):
            flush_paragraph()
            table_buffer.append(line)
            continue

        if not line.strip():
            flush_paragraph()
            continue

        if line.strip() == "---":
            flush_paragraph()
            continue

        heading = re.match(r"^(#{1,6})\s+(.+)$", line)
        if heading:
            flush_paragraph()
            level = min(len(heading.group(1)), 4)
            doc.add_heading(clean_inline(heading.group(2)), level=level)
            continue

        bullet = re.match(r"^\s*-\s+(.+)$", line)
        if bullet:
            flush_paragraph()
            doc.add_paragraph(clean_inline(bullet.group(1)), style="List Bullet")
            continue

        paragraph_buffer.append(line)

    flush_table()
    flush_paragraph()
    return doc


def main() -> None:
    markdown = SOURCE.read_text(encoding="utf-8")
    doc = render_docx(markdown)
    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
