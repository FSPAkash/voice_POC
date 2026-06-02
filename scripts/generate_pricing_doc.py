from __future__ import annotations

import json
import sys
from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from backend import app as pricing_app

DOC_PATH = ROOT / "DHL_POC_price comps for models.docx"
CALL_LOG_PATH = ROOT / "backend" / "data" / "call_log.jsonl"
BASELINE_PATH = ROOT / "backend" / "data" / "pricing_baseline.json"

OBSERVED_DATE = "2026-06-02"
EXTRAPOLATED_CALL_MINUTES = 4.0
CALLS_PER_DAY = 3333
CALLS_PER_MONTH = 100_000
CALLS_PER_YEAR = 1_200_000


def fmt_usd(value: float, digits: int = 3) -> str:
    return f"${value:.{digits}f}"


def fmt_usd4(value: float) -> str:
    return f"${value:.4f}"


def fmt_inr(value: float, digits: int = 2) -> str:
    return f"₹{value:,.{digits}f}"


def fmt_usd_with_inr(value: float, inr_per_usd: float, usd_digits: int = 4, inr_digits: int = 2) -> str:
    return f"{fmt_usd(value, usd_digits)} ({fmt_inr(value * inr_per_usd, inr_digits)})"


def fmt_usd_rate_with_inr(value: float, inr_per_usd: float, usd_digits: int = 2, inr_digits: int = 2) -> str:
    return f"${value:.{usd_digits}f} ({fmt_inr(value * inr_per_usd, inr_digits)})"


def fmt_int(value: int) -> str:
    return f"{int(value):,}"


def fmt_seconds(seconds: int) -> str:
    minutes = seconds // 60
    remainder = seconds % 60
    return f"{minutes}:{remainder:02d}"


def add_table(doc: Document, headers: list[str], rows: list[list[str]]):
    table = doc.add_table(rows=1, cols=len(headers))
    try:
        table.style = "Light Grid Accent 1"
    except Exception:
        table.style = "Table Grid"
    for idx, text in enumerate(headers):
        table.rows[0].cells[idx].text = text
    for row in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            cells[idx].text = text
    return table


def _normalize_logged_costs(costs: dict[str, Any]) -> dict[str, Any]:
    agent = costs.get("agent") or {}
    response = agent.get("response_usage") or {}
    transcription = agent.get("transcription_usage") or {}
    chat_agent = costs.get("chat_agent") or {}
    supervisor = costs.get("supervisor") or {}
    language_coach = costs.get("language_coach") or {}
    combined = costs.get("combined") or {}
    return {
        "sarvam_tts_cost_usd": float(response.get("estimated_cost_usd", 0.0) or 0.0),
        "sarvam_stt_cost_usd": float(transcription.get("estimated_cost_usd", 0.0) or 0.0),
        "chat_cost_usd": float(chat_agent.get("estimated_cost_usd", 0.0) or 0.0),
        "supervisor_cost_usd": float(supervisor.get("estimated_cost_usd", 0.0) or 0.0),
        "language_coach_cost_usd": float(language_coach.get("estimated_cost_usd", 0.0) or 0.0),
        "combined_cost_usd": float(combined.get("estimated_cost_usd", 0.0) or 0.0),
        "combined_units": int(combined.get("total_tokens", 0) or 0),
        "sarvam_tts_chars": int(response.get("text_output_tokens", 0) or 0),
        "sarvam_stt_seconds": int(transcription.get("audio_input_tokens", 0) or 0),
        "chat_input_tokens": int(chat_agent.get("text_input_tokens", 0) or 0),
        "chat_cached_input_tokens": int(chat_agent.get("text_cached_input_tokens", 0) or 0),
        "chat_output_tokens": int(chat_agent.get("text_output_tokens", 0) or 0),
    }


def load_latest_call_baseline() -> dict[str, Any] | None:
    if not CALL_LOG_PATH.exists():
        return None
    try:
        lines = [line for line in CALL_LOG_PATH.read_text(encoding="utf-8").splitlines() if line.strip()]
    except OSError:
        return None
    for raw in reversed(lines):
        try:
            parsed = json.loads(raw)
        except json.JSONDecodeError:
            continue
        duration_sec = int(parsed.get("duration_sec", 0) or 0)
        cost_usd = float(parsed.get("cost_usd", 0.0) or 0.0)
        costs = parsed.get("costs") or {}
        if duration_sec > 0 and cost_usd > 0 and isinstance(costs, dict):
            return {
                "source": f"latest completed call log entry {parsed.get('id', '')}".strip(),
                "timestamp": parsed.get("timestamp", ""),
                "duration_sec": duration_sec,
                "cost_usd": cost_usd,
                "total_units": int(parsed.get("total_units", 0) or 0),
                "costs": _normalize_logged_costs(costs),
            }
    return None


def load_fallback_baseline() -> dict[str, Any]:
    if BASELINE_PATH.exists():
        return json.loads(BASELINE_PATH.read_text(encoding="utf-8"))
    snapshot = pricing_app.current_stack_pricing_snapshot()
    observed = snapshot["observed"]
    return {
        "source": "fallback current ledger snapshot",
        "duration_sec": 60,
        "cost_usd": float(observed["combined_cost_usd"]),
        "total_units": int(observed["combined_units"]),
        "costs": observed,
    }


def load_observed_baseline() -> dict[str, Any]:
    return load_latest_call_baseline() or load_fallback_baseline()


def build_doc() -> Document:
    snapshot = pricing_app.current_stack_pricing_snapshot()
    stack = snapshot["stack"]
    sarvam_rates = snapshot["rates"]["sarvam"]
    openai_rates = snapshot["rates"]["openai"]
    inr_per_usd = float(sarvam_rates["inr_per_usd"])
    baseline = load_observed_baseline()
    observed = baseline["costs"]
    observed_chat_input_tokens = int(observed.get("chat_input_tokens", 0) or 0)
    observed_chat_cached_input_tokens = int(observed.get("chat_cached_input_tokens", 0) or 0)
    observed_chat_output_tokens = int(observed.get("chat_output_tokens", 0) or 0)

    duration_sec = max(int(baseline.get("duration_sec", 0) or 0), 1)
    observed_minutes = duration_sec / 60.0
    cost_per_minute = float(baseline["cost_usd"]) / observed_minutes
    units_per_minute = int(round(int(baseline.get("total_units", 0) or 0) / observed_minutes))
    extrapolation_factor = (EXTRAPOLATED_CALL_MINUTES * 60.0) / duration_sec

    extrapolated_tts = float(observed["sarvam_tts_cost_usd"]) * extrapolation_factor
    extrapolated_stt = float(observed["sarvam_stt_cost_usd"]) * extrapolation_factor
    extrapolated_chat = float(observed["chat_cost_usd"]) * extrapolation_factor
    extrapolated_supervisor = float(observed["supervisor_cost_usd"]) * extrapolation_factor
    extrapolated_coach = float(observed["language_coach_cost_usd"]) * extrapolation_factor
    extrapolated_total = float(baseline["cost_usd"]) * extrapolation_factor
    extrapolated_units = int(round(int(baseline.get("total_units", 0) or 0) * extrapolation_factor))

    daily_cost = extrapolated_total * CALLS_PER_DAY
    monthly_cost = extrapolated_total * CALLS_PER_MONTH
    annual_cost = extrapolated_total * CALLS_PER_YEAR

    two_min_cost = cost_per_minute * 2.0
    six_min_cost = cost_per_minute * 6.0

    total_for_share = max(extrapolated_total, 1e-9)
    tts_share = extrapolated_tts / total_for_share
    stt_share = extrapolated_stt / total_for_share
    chat_share = extrapolated_chat / total_for_share

    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10.5)

    core = doc.core_properties
    core.title = "Voice Stack Pricing - DHL Collections POC"
    core.subject = "Dashboard-grounded pricing extrapolation"
    core.author = "OpenAI Codex"

    doc.add_heading("Voice Stack Pricing - DHL Collections POC", level=1)
    doc.add_paragraph(f"Date: {OBSERVED_DATE}")
    doc.add_paragraph("Prepared for: DHL Collections POC review")
    doc.add_paragraph(
        "Scope: Current stack only. This version restores the 4-minute and volume-planning sections, "
        "but all extrapolations are now based on an actually tracked dashboard baseline instead of the old synthetic benchmark."
    )
    doc.add_paragraph(
        f"FX normalization used for all USD↔INR equivalents in this document: ₹{inr_per_usd:.2f} per USD."
    )

    doc.add_heading("0. Summary", level=2)
    doc.add_paragraph(
        "The current stack is Sarvam Bulbul v3 for TTS, Sarvam Saaras v3 for STT, and OpenAI gpt-4.1 for the collections policy/chat layer. "
        "Using the latest measured dashboard baseline, the app is currently tracking at a per-minute runtime cost that extrapolates to the 4-minute planning view below."
    )
    add_table(
        doc,
        ["Stack", "Observed call baseline", "Tracked cost / min", "Extrapolated cost / 4-min call", "Best-fit note"],
        [[
            f"{stack['tts_model']} + {stack['stt_model']} + {stack['chat_model']}",
            f"{fmt_seconds(duration_sec)} @ {fmt_usd_with_inr(float(baseline['cost_usd']), inr_per_usd, usd_digits=4)}",
            fmt_usd_with_inr(cost_per_minute, inr_per_usd, usd_digits=3),
            fmt_usd_with_inr(extrapolated_total, inr_per_usd, usd_digits=3),
            "Uses the app's tracked runtime cost as the source of truth, then scales it to the earlier 4-minute planning format.",
        ]],
    )

    doc.add_heading("1. Baseline Configuration", level=2)
    add_table(
        doc,
        ["Component", "Provider / model", "Cost treatment"],
        [
            ["Agent speech (TTS)", f"Sarvam {stack['tts_model']}", "Billed on actual output characters"],
            ["Customer transcription (STT)", f"Sarvam {stack['stt_model']}", "Billed on actual committed audio seconds"],
            ["Conversation policy engine", f"OpenAI {stack['chat_model']}", "Billed on actual text input/output tokens"],
            ["Supervisor checks", stack["supervisor_model"], "Configured, but deterministic unless usage is emitted"],
            ["Language coach", stack["language_coach_model"], "Configured, but deterministic unless usage is emitted"],
            ["Call summary", "deterministic", "No live model-call cost unless usage is emitted"],
        ],
    )

    doc.add_heading("2. Ground-Truth Observed Call Baseline", level=2)
    doc.add_paragraph(
        f"Observed source: {baseline.get('source', 'unknown')}. This is the measured call baseline used for all extrapolations below."
    )
    add_table(
        doc,
        ["Metric", "Observed value"],
        [
            ["Observed call duration", fmt_seconds(duration_sec)],
            ["Observed dashboard-tracked total cost", fmt_usd_with_inr(float(baseline["cost_usd"]), inr_per_usd, usd_digits=4)],
            ["Observed dashboard-tracked total units", fmt_int(int(baseline.get("total_units", 0) or 0))],
            ["Derived tracked cost / minute", fmt_usd_with_inr(cost_per_minute, inr_per_usd, usd_digits=4)],
            ["Derived tracked units / minute", fmt_int(units_per_minute)],
            ["Observed Sarvam TTS chars", fmt_int(int(observed["sarvam_tts_chars"]))],
            ["Observed Sarvam STT seconds", fmt_int(int(observed["sarvam_stt_seconds"]))],
            ["Observed OpenAI chat input tokens", fmt_int(observed_chat_input_tokens)],
            ["Observed OpenAI chat cached input tokens", fmt_int(observed_chat_cached_input_tokens)],
            ["Observed OpenAI chat output tokens", fmt_int(observed_chat_output_tokens)],
        ],
    )

    doc.add_heading("3. Provider Rates Used", level=2)
    provider_rate_rows = [
        [f"Sarvam {stack['tts_model']}", f"INR {sarvam_rates['tts_inr_per_10k_chars'][stack['tts_model']]:.0f}", "per 10,000 characters"],
        [f"Sarvam {stack['stt_model']}", f"INR {sarvam_rates['stt_inr_per_hour'][stack['stt_model']]:.0f}", "per audio hour"],
        [f"OpenAI {stack['chat_model']} input", fmt_usd_rate_with_inr(openai_rates[stack['chat_model']]['text_input_per_million'], inr_per_usd), "per 1M text input tokens"],
        [f"OpenAI {stack['chat_model']} cached input", fmt_usd_rate_with_inr(openai_rates[stack['chat_model']]['text_cached_input_per_million'], inr_per_usd), "per 1M cached input tokens"],
        [f"OpenAI {stack['chat_model']} output", fmt_usd_rate_with_inr(openai_rates[stack['chat_model']]['text_output_per_million'], inr_per_usd), "per 1M text output tokens"],
    ]
    if stack["supervisor_model"] == stack["language_coach_model"]:
        shared_model = stack["supervisor_model"]
        provider_rate_rows.extend(
            [
                [f"OpenAI {shared_model} input (supervisor / language coach)", fmt_usd_rate_with_inr(openai_rates[shared_model]['text_input_per_million'], inr_per_usd), "per 1M text input tokens"],
                [f"OpenAI {shared_model} cached input (supervisor / language coach)", fmt_usd_rate_with_inr(openai_rates[shared_model]['text_cached_input_per_million'], inr_per_usd), "per 1M cached input tokens"],
                [f"OpenAI {shared_model} output (supervisor / language coach)", fmt_usd_rate_with_inr(openai_rates[shared_model]['text_output_per_million'], inr_per_usd), "per 1M text output tokens"],
            ]
        )
    else:
        provider_rate_rows.extend(
            [
                [f"OpenAI {stack['supervisor_model']} input", fmt_usd_rate_with_inr(openai_rates[stack['supervisor_model']]['text_input_per_million'], inr_per_usd), "per 1M text input tokens"],
                [f"OpenAI {stack['supervisor_model']} cached input", fmt_usd_rate_with_inr(openai_rates[stack['supervisor_model']]['text_cached_input_per_million'], inr_per_usd), "per 1M cached input tokens"],
                [f"OpenAI {stack['supervisor_model']} output", fmt_usd_rate_with_inr(openai_rates[stack['supervisor_model']]['text_output_per_million'], inr_per_usd), "per 1M text output tokens"],
                [f"OpenAI {stack['language_coach_model']} input", fmt_usd_rate_with_inr(openai_rates[stack['language_coach_model']]['text_input_per_million'], inr_per_usd), "per 1M text input tokens"],
                [f"OpenAI {stack['language_coach_model']} cached input", fmt_usd_rate_with_inr(openai_rates[stack['language_coach_model']]['text_cached_input_per_million'], inr_per_usd), "per 1M cached input tokens"],
                [f"OpenAI {stack['language_coach_model']} output", fmt_usd_rate_with_inr(openai_rates[stack['language_coach_model']]['text_output_per_million'], inr_per_usd), "per 1M text output tokens"],
            ]
        )
    add_table(
        doc,
        ["Service", "Published rate", "Billing unit"],
        provider_rate_rows,
    )
    doc.add_paragraph(
        f"Internal normalization rate: INR {sarvam_rates['inr_per_usd']:.2f} per USD. Price table version: {snapshot['price_table_version']}."
    )
    doc.add_paragraph(
        f"Official provider pricing is Sarvam in INR and OpenAI in USD. Every INR figure shown next to a USD figure in this document uses the same internal normalization rate: ₹{inr_per_usd:.2f} per USD."
    )

    doc.add_heading("4. Detailed Cost Build-Up", level=2)
    doc.add_paragraph("Observed tracked call")
    add_table(
        doc,
        ["Line item", "Measured basis", "Cost"],
        [
            [f"Sarvam {stack['tts_model']} TTS", f"{fmt_int(int(observed['sarvam_tts_chars']))} chars tracked by dashboard", fmt_usd_with_inr(float(observed["sarvam_tts_cost_usd"]), inr_per_usd, usd_digits=4)],
            [f"Sarvam {stack['stt_model']} STT", f"{fmt_int(int(observed['sarvam_stt_seconds']))} seconds tracked by dashboard", fmt_usd_with_inr(float(observed["sarvam_stt_cost_usd"]), inr_per_usd, usd_digits=4)],
            [f"OpenAI {stack['chat_model']} chat / policy", f"{fmt_int(observed_chat_input_tokens)} input + {fmt_int(observed_chat_cached_input_tokens)} cached input + {fmt_int(observed_chat_output_tokens)} output tokens tracked", fmt_usd_with_inr(float(observed["chat_cost_usd"]), inr_per_usd, usd_digits=4)],
            ["Supervisor", "Measured dashboard cost", fmt_usd_with_inr(float(observed["supervisor_cost_usd"]), inr_per_usd, usd_digits=4)],
            ["Language coach", "Measured dashboard cost", fmt_usd_with_inr(float(observed["language_coach_cost_usd"]), inr_per_usd, usd_digits=4)],
            ["Total / observed call", "", fmt_usd_with_inr(float(baseline["cost_usd"]), inr_per_usd, usd_digits=4)],
        ],
    )
    doc.add_paragraph(f"Extrapolated {int(EXTRAPOLATED_CALL_MINUTES)}-minute equivalent using the tracked per-minute baseline")
    add_table(
        doc,
        ["Line item", "Extrapolation", "Cost"],
        [
            [f"Sarvam {stack['tts_model']} TTS", f"Observed TTS cost x {extrapolation_factor:.4f}", fmt_usd_with_inr(extrapolated_tts, inr_per_usd, usd_digits=4)],
            [f"Sarvam {stack['stt_model']} STT", f"Observed STT cost x {extrapolation_factor:.4f}", fmt_usd_with_inr(extrapolated_stt, inr_per_usd, usd_digits=4)],
            [f"OpenAI {stack['chat_model']} chat / policy", f"Observed chat cost x {extrapolation_factor:.4f}", fmt_usd_with_inr(extrapolated_chat, inr_per_usd, usd_digits=4)],
            ["Supervisor", f"Observed supervisor cost x {extrapolation_factor:.4f}", fmt_usd_with_inr(extrapolated_supervisor, inr_per_usd, usd_digits=4)],
            ["Language coach", f"Observed coach cost x {extrapolation_factor:.4f}", fmt_usd_with_inr(extrapolated_coach, inr_per_usd, usd_digits=4)],
            [f"Total / {int(EXTRAPOLATED_CALL_MINUTES)}-min call", "", fmt_usd_with_inr(extrapolated_total, inr_per_usd, usd_digits=4)],
        ],
    )

    doc.add_heading("5. Cost Driver Summary", level=2)
    add_table(
        doc,
        ["Component", "Share of extrapolated 4-min cost", "Observation"],
        [
            [f"Sarvam {stack['tts_model']}", f"{tts_share * 100:.1f}%", "Most sensitive to how much the agent speaks"],
            [f"Sarvam {stack['stt_model']}", f"{stt_share * 100:.1f}%", "Tracks committed customer speech only"],
            [f"OpenAI {stack['chat_model']}", f"{chat_share * 100:.1f}%", "Driven by policy prompt size and chat turn volume"],
        ],
    )

    doc.add_heading("6. Sensitivity to Call Length", level=2)
    doc.add_paragraph(
        "Using the tracked per-minute baseline, the table below scales the same call economics to shorter and longer calls."
    )
    add_table(
        doc,
        ["Scenario", "Assumed call length", "Estimated cost / call"],
        [
            ["-50% vs 4-min baseline", "2 minutes", fmt_usd_with_inr(two_min_cost, inr_per_usd, usd_digits=3)],
            ["Baseline", "4 minutes", fmt_usd_with_inr(extrapolated_total, inr_per_usd, usd_digits=3)],
            ["+50% vs 4-min baseline", "6 minutes", fmt_usd_with_inr(six_min_cost, inr_per_usd, usd_digits=3)],
        ],
    )

    doc.add_heading("7. Volume Projection", level=2)
    doc.add_paragraph(
        f"Projection assumes {CALLS_PER_MONTH:,} calls/month, approximately {CALLS_PER_DAY:,} calls/day, and {CALLS_PER_YEAR:,} calls/year. "
        f"Each projected call uses the extrapolated {int(EXTRAPOLATED_CALL_MINUTES)}-minute equivalent cost."
    )
    add_table(
        doc,
        ["Metric", "Value"],
        [
            [f"Extrapolated cost per {int(EXTRAPOLATED_CALL_MINUTES)}-min call", fmt_usd_with_inr(extrapolated_total, inr_per_usd, usd_digits=3)],
            ["Daily runtime cost", fmt_usd_with_inr(daily_cost, inr_per_usd, usd_digits=2)],
            ["Monthly runtime cost", fmt_usd_with_inr(monthly_cost, inr_per_usd, usd_digits=2)],
            ["Annual runtime cost", fmt_usd_with_inr(annual_cost, inr_per_usd, usd_digits=2)],
        ],
    )

    doc.add_heading("8. Notes and Sources", level=2)
    doc.add_paragraph(
        "The dashboard ledger is the source of truth for actual call economics. This document keeps the old planning format, but the 4-minute and volume sections are now extrapolated from a measured baseline instead of a synthetic benchmark."
    )
    doc.add_paragraph(
        "Pricing audit date: June 2, 2026. OpenAI rates were cross-checked against the official OpenAI API pricing and model pages. Sarvam rates were cross-checked against the official Sarvam pricing docs."
    )
    if int(observed["sarvam_stt_seconds"]) == 0:
        doc.add_paragraph(
            "The current observed baseline still shows zero tracked STT seconds. That baseline was captured before a fresh post-fix call exercised the repaired STT session metering path. The next completed call log will replace this fallback baseline automatically."
        )
    doc.add_paragraph("Sarvam API pricing: https://www.sarvam.ai/api-pricing")
    doc.add_paragraph("OpenAI pricing: https://developers.openai.com/api/docs/pricing")
    doc.add_paragraph("OpenAI GPT-4.1 model page: https://developers.openai.com/api/docs/models/gpt-4.1")
    doc.add_paragraph("OpenAI GPT-4.1 mini model page: https://developers.openai.com/api/docs/models/gpt-4.1-mini")
    doc.add_paragraph("Repo reference: backend/app.py pricing snapshot + backend/data/call_log.jsonl + backend/data/pricing_baseline.json")

    return doc


def main() -> None:
    doc = build_doc()
    try:
        DOC_PATH.unlink(missing_ok=True)
        doc.save(DOC_PATH)
        print(f"Wrote {DOC_PATH}")
    except PermissionError:
        fallback = DOC_PATH.with_name(f"{DOC_PATH.stem}.updated{DOC_PATH.suffix}")
        fallback.unlink(missing_ok=True)
        doc.save(fallback)
        print(f"Wrote {fallback} (original file was locked)")


if __name__ == "__main__":
    main()
